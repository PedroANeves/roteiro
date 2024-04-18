from datetime import timedelta

import pytest
from docx import Document  # type: ignore
from pytest import MonkeyPatch

from src.roteiro import (
    cleanup_line,
    cli,
    extract_lines,
    extract_timestamp_or_none,
    format_line,
    format_timedelta,
    format_timestamp,
    get_markers,
    strip_accents,
)

# false positive workaround for W0621:redefined-outer-name :
# https://github.com/pylint-dev/pylint/issues/6531
# pylint: disable=redefined-outer-name


@pytest.fixture
def sample_doc(tmp_path):
    filename = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("0000	Description	0100")
    doc.add_paragraph("Not a line")
    doc.add_paragraph(
        "0130	A Larger Description with symbols such as … or , or ?!@#		0200"
    )
    doc.add_paragraph("0130	Description with accents çáã	0200")
    doc.add_paragraph("0300	No Final timestamp")
    doc.add_paragraph("")
    doc.add_paragraph("Another Not a line after a empty line")
    doc.add_paragraph("")
    doc.save(filename)
    return filename


def test_extract_lines(sample_doc):
    assert extract_lines(sample_doc) == [
        "0000	Description	0100",
        "Not a line",
        "0130	A Larger Description with symbols such as … or , or ?!@#		0200",
        "0130	Description with accents caa	0200",
        "0300	No Final timestamp",
        "",
        "Another Not a line after a empty line",
        "",
    ]


def test_cleanup_line():
    lines = "\n0130	Description with symbols … or , or ?!@#	0200"
    assert cleanup_line(lines) == (
        "0130	Description with symbols … or , or ?!@#	0200"
    )


def test_strip_accents():
    assert strip_accents("çáãé") == "caae"


def test_extract_timestamp():
    assert extract_timestamp_or_none("0848	Description	1006") == (
        "0848",
        "Description",
        "1006",
    )


def test_extract_timestamp_optional_duration():
    assert extract_timestamp_or_none("0848	Description") == (
        "0848",
        "Description",
        None,
    )


def test_extract_timestamp_no_valid_timestamp():
    assert extract_timestamp_or_none("Not a valid timestamp") is None


def test_format_timestamp():
    assert format_timestamp("0848", "Description", "1006") == (
        "0848",  # Name
        timedelta(hours=0, minutes=8, seconds=48),  # Start
        timedelta(hours=0, minutes=1, seconds=18),  # Duration
        "Description",  # Description
    )


def test_format_timedelta():
    delta = timedelta(hours=0, minutes=8, seconds=48)
    assert format_timedelta(delta) == "08:48.000"


def test_format_line():
    assert (
        format_line("0848", "08:48.000", "01:18.000", "Description")
        == "0848	08:48.000	01:18.000	decimal	Subclip	Description"
    )


def test_get_markers(sample_doc):
    lines = get_markers(sample_doc)
    assert lines == [
        "Name\tStart\tDuration Time\tFormat\tType\tDescription",
        "0000	00:00.000	01:00.000	decimal	Subclip	Description",
        (
            "0130	01:30.000	00:30.000	decimal	Subclip	"
            "A Larger Description with symbols such as … or , or ?!@#	"
        ),
        (
            "0130	01:30.000	00:30.000	decimal	Subclip	"
            "Description with accents caa"
        ),
        "0300	03:00.000	00:10.000	decimal	Subclip	No Final timestamp",
    ]


def test_cli_prints_table(capsys, monkeypatch: MonkeyPatch):
    user_inputs = ["fake_docx", ""]
    monkeypatch.setattr("builtins.input", lambda _: user_inputs.pop(0))
    monkeypatch.setattr("os.path.isfile", lambda _: True)

    cli(lambda _filename: ["Fake extracted line"])

    out, _ = capsys.readouterr()
    assert out == "Roteiro Extractor\nFake extracted line\n"


def test_cli_prints_table_accents(capsys, monkeypatch: MonkeyPatch, sample_doc):
    user_inputs = [str(sample_doc), ""]
    monkeypatch.setattr("builtins.input", lambda _: user_inputs.pop(0))

    cli(get_markers)

    out, _ = capsys.readouterr()
    assert out == (
        "Roteiro Extractor\n"
        "Name\tStart\tDuration Time\tFormat\tType\tDescription\n"
        "0000	00:00.000	01:00.000	decimal	Subclip	Description\n"
        "0130	01:30.000	00:30.000	"
        "decimal	Subclip	A Larger Description with symbols such as "
        "… or , or ?!@#	\n"
        "0130	01:30.000	00:30.000	"
        "decimal	Subclip	Description with accents caa\n"
        "0300	03:00.000	00:10.000	decimal	Subclip	No Final timestamp\n"
    )


def test_accept_paragraphs_starting_with_newline(tmp_path):
    filename = tmp_path / "with_newline.docx"
    doc = Document()
    doc.add_paragraph("\n0445\tDescription with starting newline\t0511")
    doc.save(filename)

    assert "0445\tDescription with starting newline\t0511" in extract_lines(
        filename
    )


def test_sometimes_spaces_separates_columns_instead_of_tabs():
    assert extract_timestamp_or_none("1611 Description with space\t1721") == (
        "1611",
        "Description with space",
        "1721",
    )
